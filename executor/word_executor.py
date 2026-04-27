鬠ｿ# modules/word_executor.py
import logging
import os

logger = logging.getLogger("OfficeAgent")


def _wd_color(hex_color):
    from docx.shared import RGBColor
    named = {
        "red": "FF0000",
        "green": "00B050",
        "blue": "0070C0",
        "yellow": "FFFF00",
        "orange": "FFA500",
        "purple": "7030A0",
        "pink": "FF69B4",
        "black": "000000",
        "white": "FFFFFF",
        "gray": "808080",
        "grey": "808080",
        "navy": "000080",
    }
    raw = str(hex_color or "000000").strip().lower()
    h = named.get(raw, raw).lstrip("#")
    if len(h) == 8:
        h = h[-6:]
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


class WordExecutor:
    def __init__(self, doc):
        self.doc = doc

    def run(self, action_dict):
        action  = action_dict.get("action")
        handler = getattr(self, f"_do_{action}", None)
        if not handler:
            logger.warning(f"Word: Unknown action '{action}'")
            return False
        try:
            handler(action_dict)
            return True
        except Exception as e:
            logger.error(f"Word action '{action}' failed: {e}")
            return False

    def _targeted_paragraphs(self, p):
        target = str(p.get("target", "selection") or "selection").strip()
        paragraphs = list(self.doc.paragraphs)
        if not paragraphs:
            return []
        if target.lower() == "selection":
            return paragraphs

        target_lower = target.lower()
        matched = [para for para in paragraphs if target_lower in para.text.lower()]
        return matched or paragraphs

    def _targeted_runs(self, p):
        runs = []
        for para in self._targeted_paragraphs(p):
            runs.extend(list(para.runs))
        return runs

    def _highlight_color(self, raw_color):
        from docx.enum.text import WD_COLOR_INDEX

        color = str(raw_color or "yellow").strip().lower()
        hex_map = {
            "ffff00": "yellow",
            "00b050": "green",
            "00ff00": "green",
            "00ffff": "cyan",
            "ff69b4": "pink",
            "ff0000": "red",
            "0070c0": "blue",
            "0000ff": "blue",
            "808080": "gray",
            "808080ff": "gray",
        }
        color = hex_map.get(color.lstrip("#"), color)
        color_map = {
            "yellow": WD_COLOR_INDEX.YELLOW,
            "green": WD_COLOR_INDEX.BRIGHT_GREEN,
            "cyan": WD_COLOR_INDEX.CYAN,
            "pink": WD_COLOR_INDEX.PINK,
            "red": WD_COLOR_INDEX.RED,
            "blue": WD_COLOR_INDEX.BLUE,
            "gray": WD_COLOR_INDEX.GRAY_25,
            "grey": WD_COLOR_INDEX.GRAY_25,
        }
        return color_map.get(color, WD_COLOR_INDEX.YELLOW)

    def _do_create_document(self, p):
        from docx import Document
        self.doc = Document()

    def _do_open_document(self, p):
        # Open/load is handled by server-level lifecycle.
        return

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Document Content ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_add_paragraph(self, p):
        self.doc.add_paragraph(p.get("text", ""))

    def _do_add_heading(self, p):
        self.doc.add_heading(p.get("text", ""), level=int(p.get("level", 1)))

    def _do_add_table(self, p):
        self.doc.add_table(rows=int(p.get("rows", 2)), cols=int(p.get("cols", 2)))

    def _do_delete_table(self, p):
        if self.doc.tables:
            tbl = self.doc.tables[-1]._element
            tbl.getparent().remove(tbl)

    def _do_add_table_row(self, p):
        if self.doc.tables:
            self.doc.tables[-1].add_row()

    def _do_add_table_column(self, p):
        logger.info("Add table column (requires win32com for full support)")

    def _do_set_table_style(self, p):
        if self.doc.tables:
            try:
                self.doc.tables[-1].style = p.get("style", "Table Grid")
            except Exception:
                pass

    def _do_add_bullet_list(self, p):
        for item in p.get("items", []):
            self.doc.add_paragraph(str(item), style="List Bullet")

    def _do_add_numbered_list(self, p):
        for item in p.get("items", []):
            self.doc.add_paragraph(str(item), style="List Number")

    def _do_continue_list(self, p):
        self.doc.add_paragraph(p.get("text", ""), style="List Bullet")

    def _do_remove_list_format(self, p):
        for para in self.doc.paragraphs:
            if para.style.name.startswith("List"):
                para.style = self.doc.styles["Normal"]

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Font & Style ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_set_bold(self, p):
        for run in self._targeted_runs(p):
            run.bold = p.get("bold", True)

    def _do_set_italic(self, p):
        for run in self._targeted_runs(p):
            run.italic = p.get("italic", True)

    def _do_set_underline(self, p):
        for run in self._targeted_runs(p):
            run.underline = p.get("underline", True)

    def _do_remove_underline(self, p):
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.underline = False

    def _do_set_strikethrough(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        for run in self._targeted_runs(p):
            rPr = run._r.get_or_add_rPr()
            strike = etree.SubElement(rPr, qn("w:strike"))
            strike.set(qn("w:val"), "true")

    def _do_remove_strikethrough(self, p):
        from docx.oxml.ns import qn
        for run in self._targeted_runs(p):
            rPr = run._r.get_or_add_rPr()
            strike = rPr.find(qn("w:strike"))
            if strike is not None:
                rPr.remove(strike)

    def _do_set_superscript(self, p):
        for run in self._targeted_runs(p):
            run.font.superscript = True

    def _do_set_subscript(self, p):
        for run in self._targeted_runs(p):
            run.font.subscript = True

    def _do_set_font_size(self, p):
        from docx.shared import Pt
        for run in self._targeted_runs(p):
            run.font.size = Pt(int(p["size"]))

    def _do_set_font_name(self, p):
        for run in self._targeted_runs(p):
            run.font.name = p["name"]

    def _do_set_font_color(self, p):
        for run in self._targeted_runs(p):
            run.font.color.rgb = _wd_color(p["color"])

    def _do_set_highlight(self, p):
        color = self._highlight_color(p.get("color", "yellow"))
        for run in self._targeted_runs(p):
            run.font.highlight_color = color

    def _do_remove_highlight(self, p):
        from docx.enum.text import WD_COLOR_INDEX
        for run in self._targeted_runs(p):
            run.font.highlight_color = WD_COLOR_INDEX.AUTO

    def _do_change_case(self, p):
        case = p.get("case", "upper")
        for run in self._targeted_runs(p):
            if case == "upper":
                run.text = run.text.upper()
            elif case == "lower":
                run.text = run.text.lower()
            elif case == "title":
                run.text = run.text.title()

    def _do_clear_formatting(self, p):
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.bold      = None
                run.italic    = None
                run.underline = None
                run.font.size = None
                run.font.name = None

    def _do_apply_style(self, p):
        for para in self._targeted_paragraphs(p):
            try:
                para.style = p.get("style", "Normal")
            except Exception:
                pass

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Alignment & Spacing ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_set_alignment(self, p):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            "center":  WD_ALIGN_PARAGRAPH.CENTER,
            "right":   WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left":    WD_ALIGN_PARAGRAPH.LEFT,
        }
        align = align_map.get(p.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        for para in self._targeted_paragraphs(p):
            para.alignment = align

    def _do_set_line_spacing(self, p):
        spacing = float(p.get("spacing", 1.15))
        for para in self._targeted_paragraphs(p):
            para.paragraph_format.line_spacing = spacing

    def _do_set_paragraph_spacing(self, p):
        from docx.shared import Pt
        for para in self._targeted_paragraphs(p):
            para.paragraph_format.space_before = Pt(int(p.get("before", 0)))
            para.paragraph_format.space_after  = Pt(int(p.get("after",  8)))

    def _do_set_indent(self, p):
        from docx.shared import Inches
        for para in self._targeted_paragraphs(p):
            para.paragraph_format.first_line_indent = Inches(float(p.get("indent", 0.5)))

    def _do_remove_indent(self, p):
        for para in self._targeted_paragraphs(p):
            para.paragraph_format.first_line_indent = None

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Page Layout ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_set_margins(self, p):
        from docx.shared import Inches
        section               = self.doc.sections[0]
        section.top_margin    = Inches(float(p.get("top",    1)))
        section.bottom_margin = Inches(float(p.get("bottom", 1)))
        section.left_margin   = Inches(float(p.get("left",   1)))
        section.right_margin  = Inches(float(p.get("right",  1)))

    def _do_set_orientation(self, p):
        from docx.enum.section import WD_ORIENT
        section = self.doc.sections[0]
        if p.get("orientation") == "landscape":
            section.orientation             = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

    def _do_set_paper_size(self, p):
        from docx.shared import Mm, Inches
        section = self.doc.sections[0]
        size    = p.get("size", "").lower()
        if size == "a4":
            section.page_width  = Mm(210)
            section.page_height = Mm(297)
        elif size == "a3":
            section.page_width  = Mm(297)
            section.page_height = Mm(420)
        elif size in ("letter", "us"):
            section.page_width  = Inches(8.5)
            section.page_height = Inches(11)

    def _do_insert_page_break(self, p):
        from docx.enum.text import WD_BREAK
        para = self.doc.add_paragraph()
        run  = para.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _do_insert_section_break(self, p):
        self.doc.add_section()

    def _do_set_columns(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        count   = int(p.get("count", 2))
        sect_pr = self.doc.sections[0]._sectPr
        cols    = etree.SubElement(sect_pr, qn("w:cols"))
        cols.set(qn("w:num"), str(count))

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Header / Footer ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_add_header(self, p):
        self.doc.sections[0].header.paragraphs[0].text = p.get("text", "")

    def _do_add_footer(self, p):
        self.doc.sections[0].footer.paragraphs[0].text = p.get("text", "")

    def _do_remove_header(self, p):
        self.doc.sections[0].header.paragraphs[0].text = ""

    def _do_remove_footer(self, p):
        self.doc.sections[0].footer.paragraphs[0].text = ""

    def _do_add_page_number(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        footer = self.doc.sections[0].footer
        para   = footer.paragraphs[0]
        run    = para.add_run()
        fld    = etree.fromstring(
            '<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:fldCharType="begin"/>'
        )
        run._r.append(fld)

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ References & Notes ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_add_table_of_contents(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        para  = self.doc.add_paragraph()
        run   = para.add_run()
        run._r.append(etree.fromstring(
            '<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:fldCharType="begin"/>'
        ))
        run2  = para.add_run()
        run2._r.append(etree.fromstring(
            '<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xml:space="preserve"> TOC ¥¥o "1-3" ¥¥h ¥¥z ¥¥u </w:instrText>'
        ))

    def _do_update_table_of_contents(self, p):
        logger.info("Update TOC (requires win32com)")

    def _do_add_footnote(self, p):
        para = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
        para.add_run(f" [{p.get('text', '')}]")

    def _do_add_endnote(self, p):
        para = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
        para.add_run(f" (end: {p.get('text', '')})")

    def _do_add_comment(self, p):
        para = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
        para.add_run(f"  /* {p.get('text', '')} */")

    def _do_delete_comment(self, p):
        logger.info("Delete comment (requires win32com)")

    def _do_add_bookmark(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        para    = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
        bm_name = p.get("name", "bookmark1")
        bm_id   = "1"
        bm_start = etree.fromstring(
            f'<w:bookmarkStart xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:id="{bm_id}" w:name="{bm_name}"/>'
        )
        para._p.append(bm_start)

    def _do_add_cross_reference(self, p):
        logger.info(f"Cross reference to '{p.get('target')}' (requires win32com)")

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Insert Elements ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_insert_image(self, p):
        from docx.shared import Inches
        path = p.get("path", "")
        if path and os.path.exists(path):
            self.doc.add_picture(path, width=Inches(4))
        else:
            logger.warning(f"Image not found: {path}")

    def _do_insert_hyperlink(self, p):
        para     = self.doc.add_paragraph()
        run      = para.add_run(p.get("text", p.get("url", "")))
        run.font.color.rgb = _wd_color("0070C0")
        run.underline      = True

    def _do_insert_horizontal_line(self, p):
        from docx.oxml.ns import qn
        from lxml import etree
        para   = self.doc.add_paragraph()
        pPr    = para._p.get_or_add_pPr()
        pBdr   = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")

    def _do_insert_special_character(self, p):
        self.doc.add_paragraph(p.get("character", ""))

    def _do_insert_date(self, p):
        from datetime import date
        self.doc.add_paragraph(str(date.today()))

    def _do_insert_text_box(self, p):
        logger.info(f"Text box with '{p.get('text')}' (requires win32com for floating boxes)")
        self.doc.add_paragraph(p.get("text", ""))

    def _do_insert_shape(self, p):
        logger.info(f"Shape '{p.get('shape_type')}' (requires win32com)")

    def _do_insert_chart(self, p):
        logger.info(f"Chart '{p.get('chart_type')}' (requires win32com)")

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Review Tools ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_find_text(self, p):
        results = [
            para.text for para in self.doc.paragraphs
            if p.get("text", "").lower() in para.text.lower()
        ]
        logger.info(f"Found in {len(results)} paragraph(s)")
        return results

    def _do_find_replace(self, p):
        find    = p.get("find_text", "")
        replace = p.get("replace_text", "")
        for para in self.doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)

    def _do_track_changes(self, p):
        logger.info(f"Track changes {'ON' if p.get('enabled') else 'OFF'} (requires win32com)")

    def _do_accept_changes(self, p):
        logger.info("Accept all changes (requires win32com)")

    def _do_reject_changes(self, p):
        logger.info("Reject all changes (requires win32com)")

    def _do_spell_check(self, p):
        logger.info("Spell check (requires win32com)")

    def _do_get_word_count(self, p):
        count = sum(len(para.text.split()) for para in self.doc.paragraphs)
        logger.info(f"Word count: {count}")
        return count

    def _do_mail_merge(self, p):
        logger.info(f"Mail merge from '{p.get('data_source')}' (requires win32com)")

    def _do_start_mail_merge(self, p):
        self._do_mail_merge(p)

    def _do_compare_documents(self, p):
        logger.info(f"Compare with '{p.get('path')}' (requires win32com)")

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Protection ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_protect_document(self, p):
        logger.info("Document protection (requires win32com)")

    def _do_unprotect_document(self, p):
        logger.info("Document unprotection (requires win32com)")

    # ﾃ｢窶昶ぎﾃ｢窶昶ぎ Save / Export ﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎﾃ｢窶昶ぎ

    def _do_save_document(self, p):
        path = getattr(self.doc, "_path", "output.docx")
        self.doc.save(path)

    def _do_save_document_as(self, p):
        filename = p.get("filename", "output")
        if not filename.endswith(".docx"):
            filename += ".docx"
        self.doc.save(filename)

    def _do_close_document(self, p):
        logger.info("Document closed (session ended)")

    def _do_export_pdf(self, p):
        logger.info(f"PDF export to '{p.get('path', 'output.pdf')}' (requires win32com)")

    def _do_print_document(self, p):
        logger.info("Print document (requires win32com)")

    def _do_undo(self, p):
        logger.info("Undo (requires win32com)")

    def _do_redo(self, p):
        logger.info("Redo (requires win32com)")
